import numpy as np
import pandas as pd
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.shared import Pt



WNS_COLS_NUM = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}num"


def set_number_of_columns(section, cols):
    """ sets number of columns through xpath. """
    section._sectPr.xpath("./w:cols")[0].set(WNS_COLS_NUM, str(cols))

def yaz(document,text,aligment,spacing,style='Normal',girinti=0):
    p = document.add_paragraph(text, style=style)
    p_format=p.paragraph_format
    p_format.alignment = aligment
    p_format.space_after = Pt(0)
    p_format.space_before = Pt(0)
    p.line_spacing = spacing
    p_format.left_indent = Inches(girinti)

def dokuman_yap(dersAdi,egitimYili,sinav,grup,anahtar=0):
    #Tanımlamalar
    okulAdi="TÜRK TELEKOM MESLEKİ VE TEKNİK ANADOLU LİSESİ"
    ogretmen="M. Başar ACAROĞLU"



    # Döküman Oluşturma ****************************
    document = Document()
    sections = document.sections
    section = sections[0]
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)

#******BAŞLIK KISMI ************************************************
    yaz(document,okulAdi+' \n  '+dersAdi.upper()+' \n '+egitimYili+' EĞİTİM ÖĞRETİM YILI '+donem+'. DÖNEM '+sinav+'. YAZILI '+grup+' GRUBU SINAV SORULARI',WD_ALIGN_PARAGRAPH.CENTER,1)
    yaz(document,'Öğrencinin Adı\t:\nNumarası\t:',WD_ALIGN_PARAGRAPH.LEFT,1)
    siklar=np.array(['','a) ','b) ', 'c) ','d) ','e) '])
#******2 SUTUN YAP****************************************************
    new_section = document.add_section(WD_SECTION.ODD_PAGE)
    new_section.start_type=WD_SECTION.CONTINUOUS
    set_number_of_columns(section, 2)


#****EXCEL DEN OKU KARIŞTIR YAZ********************************************************
    data=pd.read_excel(r'Sorular.xlsx')
    data = data.sample(frac = 1)
# df = pd.DataFrame(data,columns=['sorum'])

    dogrular=[]
    for index, row in data.iterrows():
        art=1;
        cevaplar=np.array([1,2,3,4,5])
        np.random.shuffle(cevaplar)
        yaz(document,row[0],WD_ALIGN_PARAGRAPH.LEFT,1,'List Number')

        for i in range(0,5):
            yaz(document,siklar[art]+str(row[cevaplar[i]]),WD_ALIGN_PARAGRAPH.LEFT,1,'Normal',0.1)
            if(cevaplar[i]==5):
                dogrular.append(siklar[art])
            art=art+1
    # print( row[cevaplar[0]],row[cevaplar[1]],row[cevaplar[2]],row[cevaplar[3]],row[cevaplar[4]], siklar[x])
#***CEVAP ANAHTARI*************************************************************************************************************************
    if(anahtar>0):
        table = document.add_table(rows=0, cols=6)
        table.autofit=True
        table.allow_autofit = True
        table.style = 'Table Grid'
        for i in range(1,21):
            row_cells=table.add_row().cells
            row_cells[0].text=str(i)
            cell=row_cells[0]
            cell.width=Inches(0.2)
            for j in range(1,6):
                row_cells[j].text=siklar[j][:-2]
                cell=row_cells[j]
                cell.width=Inches(0.2)
        
        

#***BAŞARILAR BÖLÜMÜ********************************************************************************************************
    yaz(document,'Başarılar Dilerim \n'+ogretmen,WD_ALIGN_PARAGRAPH.CENTER,1)


#***CEVAPLAR *****************************************************************************************************************
    document.add_page_break()
    yaz(document,'CEVAPLAR',WD_ALIGN_PARAGRAPH.CENTER,1)

    art=1;
    for x in dogrular:
        yaz(document,str(art) + ') '+ x[:-2],WD_ALIGN_PARAGRAPH.LEFT,1)
        art=art+1


    ad=dersAdi+" "+donem+". Dönem"+sinav+". Sınavı"+" "+grup+"Grubu"+".docx"
    document.save(ad)

dersAdi=input("Dersin Adını girin : ")
egitimYili='2020 - 2021'
donem=input('Kaçıncı dönem sınavı : ')
sinav=input('Dönemin kaçıncı sınavı :')


grup=["A","B"]
for x in grup:
    dokuman_yap(dersAdi.upper(),egitimYili,sinav,x)


